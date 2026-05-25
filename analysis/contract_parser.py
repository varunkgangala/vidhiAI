"""
contract_parser.py

Handles text extraction from PDF, DOCX and TXT files.

PDF extraction pipeline (tries methods in order):
  Method 1: pypdf          — fast, works for digital PDFs
  Method 2: pdfplumber     — better for complex layouts
  Method 3: pypdfium2      — Chromium PDF engine
  Method 4: OCR            — for scanned / Print-to-PDF / image-based PDFs
             Uses pdf2image + pytesseract
             Works for ALL PDF types including scanned documents

DOCX extraction:
  - Paragraphs, tables, headers via python-docx
"""

import os
import re
import pytesseract

# Windows Tesseract path - update if installed to different location
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"


# ── PDF Method 1: pypdf ───────────────────────────────────────────────────────

def _extract_pypdf(file_stream) -> str:
    try:
        from pypdf import PdfReader
        file_stream.seek(0)
        reader = PdfReader(file_stream)
        if reader.is_encrypted:
            raise ValueError(
                "This PDF is password protected. "
                "Please remove the password and try again."
            )
        text = "\n".join(
            page.extract_text() or "" for page in reader.pages
        ).strip()
        return text
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"pypdf failed: {e}")


# ── PDF Method 2: pdfplumber ──────────────────────────────────────────────────

def _extract_pdfplumber(file_stream) -> str:
    try:
        import pdfplumber
        file_stream.seek(0)
        texts = []
        with pdfplumber.open(file_stream) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texts.append(t)
        return "\n".join(texts).strip()
    except Exception as e:
        raise ValueError(f"pdfplumber failed: {e}")


# ── PDF Method 3: pypdfium2 ───────────────────────────────────────────────────

def _extract_pypdfium2(file_stream) -> str:
    try:
        import pypdfium2 as pdfium
        file_stream.seek(0)
        data = file_stream.read()
        pdf  = pdfium.PdfDocument(data)
        texts = []
        for i in range(len(pdf)):
            page     = pdf[i]
            textpage = page.get_textpage()
            t        = textpage.get_text_range()
            if t and t.strip():
                texts.append(t)
        return "\n".join(texts).strip()
    except Exception as e:
        raise ValueError(f"pypdfium2 failed: {e}")


# ── PDF Method 4: OCR ─────────────────────────────────────────────────────────

def _preprocess_image(img):
    """
    Enhance image quality for better OCR accuracy.
    Converts to grayscale and increases contrast.
    """
    try:
        from PIL import Image, ImageFilter, ImageEnhance
        # Convert to grayscale
        img = img.convert("L")
        # Sharpen
        img = img.filter(ImageFilter.SHARPEN)
        # Increase contrast
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)
        return img
    except Exception:
        return img


def _extract_ocr(file_stream) -> str:
    """
    OCR extraction — works for ALL PDF types:
      - Scanned physical documents
      - Print-to-PDF files
      - Image-based PDFs
      - Mixed PDFs

    Uses pdf2image to render pages + pytesseract to read text.
    Requires: Tesseract OCR installed on the system.
    """
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        from PIL import Image

        file_stream.seek(0)
        pdf_bytes = file_stream.read()

        print("[VidhiAI] Scanned/image PDF detected — using OCR...")
        pages = convert_from_bytes(pdf_bytes, dpi=250)
        print(f"[VidhiAI] OCR processing {len(pages)} page(s)...")

        texts = []
        for i, page_img in enumerate(pages):
            print(f"[VidhiAI] OCR page {i+1}/{len(pages)}...")
            # Preprocess for better accuracy
            processed = _preprocess_image(page_img)
            # Run OCR with page segmentation mode 6 (block of text)
            text = pytesseract.image_to_string(
                processed,
                lang="eng",
                config="--psm 6"
            )
            cleaned = _clean_ocr_text(text)
            if cleaned:
                texts.append(cleaned)

        result = "\n\n".join(texts).strip()
        print(f"[VidhiAI] OCR complete: {len(result)} characters extracted")
        return result

    except ImportError as e:
        missing = str(e).split("'")[1] if "'" in str(e) else str(e)
        raise ValueError(
            f"OCR requires '{missing}'.\n"
            "Run: pip install pytesseract pdf2image Pillow\n"
            "Also install Tesseract from: https://github.com/UB-Mannheim/tesseract/wiki"
        )
    except Exception as e:
        if "tesseract" in str(e).lower():
            raise ValueError(
                "Tesseract OCR is not installed or not found.\n\n"
                "Install it from:\n"
                "https://github.com/UB-Mannheim/tesseract/wiki\n\n"
                "Download: tesseract-ocr-w64-setup-5.x.x.exe\n"
                "Install with default settings.\n"
                "Then restart your terminal and run python app.py again."
            )
        raise ValueError(f"OCR failed: {e}")


def _clean_ocr_text(text: str) -> str:
    """Clean up common OCR artifacts."""
    if not text:
        return ""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {3,}', ' ', text)
    # Remove page numbers standing alone
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
    return text.strip()


# ── Main PDF extractor ────────────────────────────────────────────────────────

def extract_text_from_pdf(file_stream) -> str:
    """
    Smart PDF text extraction.
    Tries 4 methods automatically — no manual intervention needed.

    Works for:
      - Normal digital PDFs
      - PDFs saved from Word/Google Docs
      - Print-to-PDF files
      - Scanned documents
      - Image-based PDFs
    """
    errors  = []
    methods = [
        ("pypdf",      _extract_pypdf),
        ("pdfplumber", _extract_pdfplumber),
        ("pypdfium2",  _extract_pypdfium2),
        ("OCR",        _extract_ocr),
    ]

    for method_name, method_fn in methods:
        try:
            if hasattr(file_stream, 'seek'):
                file_stream.seek(0)

            text = method_fn(file_stream)

            if text and len(text.strip()) >= 50:
                print(f"[VidhiAI] PDF extracted via {method_name}: {len(text)} chars")
                return text
            else:
                chars = len(text.strip()) if text else 0
                msg   = f"{method_name}: only {chars} chars — trying next method"
                errors.append(msg)
                print(f"[VidhiAI] {msg}")

        except ValueError as e:
            err_msg = str(e)
            # Immediately raise password errors
            if "password" in err_msg.lower():
                raise
            # Immediately raise Tesseract installation errors
            if "tesseract" in err_msg.lower() or "install" in err_msg.lower():
                raise
            errors.append(f"{method_name}: {err_msg[:100]}")
            print(f"[VidhiAI] {method_name} failed: {err_msg[:100]}")

        except Exception as e:
            errors.append(f"{method_name}: {str(e)[:100]}")
            print(f"[VidhiAI] {method_name} error: {str(e)[:100]}")

    # All methods failed
    raise ValueError(
        "Could not extract text from this PDF after trying all methods.\n\n"
        "Methods attempted:\n" +
        "\n".join(f"  - {e}" for e in errors) +
        "\n\nBest fix options:\n"
        "1. Install Tesseract OCR (for scanned PDFs):\n"
        "   https://github.com/UB-Mannheim/tesseract/wiki\n"
        "2. Copy text from the PDF and use the 'Paste Text' tab\n"
        "3. Save as PDF from Word using File > Save As (not Print to PDF)"
    )


# ── DOCX extractor ────────────────────────────────────────────────────────────

def extract_text_from_docx(file_stream) -> str:
    """Extract text from DOCX — paragraphs, tables, and headers."""
    try:
        from docx import Document
        if hasattr(file_stream, 'seek'):
            file_stream.seek(0)
        doc   = Document(file_stream)
        texts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    texts.append(row_text)

        # Headers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        texts.append(para.text.strip())

        result = "\n".join(texts).strip()

        if len(result) < 50:
            raise ValueError(
                "DOCX file appears empty or has no readable text.\n"
                "Please check the document has content."
            )

        print(f"[VidhiAI] DOCX extracted: {len(result)} chars, "
              f"{len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        return result

    except ValueError:
        raise
    except Exception as e:
        raise ValueError(
            f"Could not read DOCX file: {e}\n"
            "Make sure it is a valid .docx format.\n"
            "Try resaving from Word as .docx and uploading again."
        )


# ── TXT extractor ─────────────────────────────────────────────────────────────

def extract_text_from_txt(file_stream) -> str:
    try:
        if hasattr(file_stream, 'seek'):
            file_stream.seek(0)
        raw  = file_stream.read()
        text = raw.decode("utf-8", errors="replace").strip()
        if len(text) < 50:
            raise ValueError("Text file is too short or empty.")
        print(f"[VidhiAI] TXT extracted: {len(text)} chars")
        return text
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"TXT extraction failed: {e}")


# ── Main parser ───────────────────────────────────────────────────────────────

def parse_contract(file, filename: str) -> str:
    """Parse contract and return extracted text. Supports PDF, DOCX, TXT."""
    ext = os.path.splitext(filename)[1].lower()
    print(f"[VidhiAI] Parsing: {filename} (type: {ext})")

    if ext == ".pdf":
        return extract_text_from_pdf(file)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file)
    elif ext == ".txt":
        return extract_text_from_txt(file)
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'.\n"
            "Please upload PDF, DOCX, or TXT."
        )


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}


def is_allowed_file(filename: str) -> bool:
    return os.path.splitext(filename)[1].lower() in ALLOWED_EXTENSIONS